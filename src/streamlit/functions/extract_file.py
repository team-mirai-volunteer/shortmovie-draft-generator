import json
import os
import shutil
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

import fitz
import pandas as pd
from docx import Document
from whoosh import index
from whoosh.fields import ID, TEXT, Schema
from whoosh.qparser import QueryParser

import streamlit as st

from src.streamlit.functions.directory import search_directory
from src.streamlit.functions.extract_ppt import extract_ppt_info, extract_text_from_ppt_file_simple


def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""

    for page in doc:
        # テーブルの抽出
        tabs = page.find_tables()

        if not tabs.tables:
            # テーブルがない場合は通常のテキスト抽出
            text += page.get_text().strip() + "\n\n"
        else:
            # テーブルがある場合は、テーブル以外のテキストを抽出し、テーブルを挿入
            page_text = ""
            last_end = 0

            for table in tabs.tables:
                # テーブル前のテキストを追加
                page_text += (
                    page.get_text(
                        "text",
                        clip=fitz.Rect(0, last_end, page.rect.width, table.bbox[1]),
                    ).strip()
                    + "\n\n"
                )

                # テーブルをマークダウン形式で追加
                df = pd.DataFrame(table.extract()[1:], columns=table.extract()[0])
                markdown_table = df.to_markdown(index=False)
                page_text += markdown_table + "\n\n"

                last_end = table.bbox[3]

            # 最後のテーブル以降のテキストを追加
            page_text += (
                page.get_text(
                    "text",
                    clip=fitz.Rect(0, last_end, page.rect.width, page.rect.height),
                ).strip()
                + "\n\n"
            )

            text += page_text

    return text.strip()


def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    table_index = 0  # テーブルのインデックスを追跡

    for element in doc.element.body:
        if element.tag.endswith("p"):  # 段落の場合
            paragraph = element.text.strip() if element.text else ""
            if paragraph:
                text += paragraph + "\n\n"

        elif element.tag.endswith("tbl"):  # テーブルの場合
            # 現在のインデックスのテーブルを取得
            table_element = doc.tables[table_index]
            rows = []

            for row in table_element.rows:
                # セル内の改行を空白に置換
                row_data = [cell.text.strip().replace("\n", "<br>") if cell.text else "" for cell in row.cells]
                rows.append(row_data)

            if rows:
                max_cols = max(len(row) for row in rows)
                padded_rows = [row + [""] * (max_cols - len(row)) for row in rows]

                if len(padded_rows) > 1:
                    df = pd.DataFrame(padded_rows[1:], columns=padded_rows[0])
                    markdown_table = df.to_markdown(index=False)
                    text += markdown_table + "\n\n"

            table_index += 1  # 次のテーブルのために増加

    return text.strip()


# ファイルの内容をテキストとして抽出する
def get_file_text(file_path: Path, simple_ppt: bool = False) -> Any:
    full_path = os.path.expanduser(file_path)
    if not os.path.exists(full_path):
        return f"指定されたファイル '{full_path}' は存在しません。"
    base_name = os.path.basename(full_path)
    file_extension = os.path.splitext(full_path)[1].lower()

    try:
        with open(full_path, "rb") as file:
            if file_extension == ".pdf":
                return extract_text_from_pdf(file)
            elif file_extension == ".docx":
                return extract_text_from_docx(file)
            elif file_extension == ".pptx":
                if simple_ppt:
                    return extract_text_from_ppt_file_simple(full_path)
                else:
                    return json.dumps(
                        extract_ppt_info(file, include_images=False),
                        ensure_ascii=False,
                        indent=2,
                    )
            elif file_extension == ".doc":
                os.system(f"antiword {full_path} > {base_name}.txt")
                with open(f"{base_name}.txt", "r", encoding="utf-8") as txt_file:
                    file_text = txt_file.read()
                os.remove(f"{base_name}.txt")
                return file_text
            elif file_extension in [".txt", ".md", ".csv", ".tsv", ".py", ".json", ".yaml", ".yml"]:
                import chardet

                try:
                    raw_data = file.read()
                    detected_encoding = chardet.detect(raw_data).get("encoding", "utf-8")
                    if not detected_encoding:
                        detected_encoding = "utf-8"
                    content = raw_data.decode(detected_encoding, errors="replace")
                    return content
                except Exception as e:
                    return f"テキストファイルの読み込みに失敗しました: {str(e)}"
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
    except Exception as e:
        return f"Failed to extract text from file: {str(e)}"


# ファイル操作関数の実装
def explore_directory(
    query: str,
    target_dir: str | list[str],
    max_depth: Optional[int] = 2,
    file_types: Optional[list[str]] = ["pdf", "pptx", "docx", "csv", "xls", "xlsx"],
) -> str:
    if isinstance(target_dir, list):
        results = []
        for dir_path in target_dir:
            result = search_directory(
                path=dir_path,
                max_depth=max_depth,
                pattern=query if query else None,
                file_types=file_types if file_types else None,
            )
            results.append(result)
        return "\n\n".join(results)
    else:
        return search_directory(
            path=target_dir,
            max_depth=max_depth,
            pattern=query if query else None,
            file_types=file_types if file_types else None,
        )


# ファイルの全文検索
def keyword_search(keywords: List[str], dir_path: Path | str, reindex: bool = False) -> str:
    """
    ディレクトリに含まれるファイルの中に、指定されたキーワードが含まれるファイルを検索して返す。
    検索結果を文字列形式で返す。
    """
    # ディレクトリパスをフルパスに変換
    full_dir_path = Path(str(os.path.expanduser(dir_path)))

    # インデックスディレクトリのベースパス
    base_index_dir = "whoosh_indices"
    if not os.path.exists(base_index_dir):
        os.makedirs(base_index_dir)

    # ディレクトリパスからインデックス名を生成（パスの文字列をハッシュ化）
    safe_dir_name = str(full_dir_path).replace("/", "_").replace("\\", "_").replace("~", "home")
    index_dir = os.path.join(base_index_dir, safe_dir_name)

    # スキーマの定義
    schema = Schema(path=ID(stored=True), content=TEXT(stored=True))

    # インデックスの作成
    if (not os.path.exists(index_dir)) or reindex:
        os.makedirs(index_dir, exist_ok=True)
        with st.spinner(f"インデックスを作成中... ({full_dir_path})"):
            ix = index.create_in(index_dir, schema)
            writer = ix.writer()

            # まず全ファイル数を数える（~$で始まるファイル、および"."で始まるディレクトリを除外）
            total_files = len([f for f in full_dir_path.rglob("*") if not os.path.basename(f).startswith("~$") and not any(part.startswith(".") for part in f.parts)])
            progress_bar = st.progress(0)

            # ファイルの走査とインデックス作成（"."で始まるディレクトリはスキップ）
            for i, file_path in enumerate([f for f in full_dir_path.rglob("*") if not any(part.startswith(".") for part in f.parts)]):
                if file_path.is_file() and not os.path.basename(file_path).startswith("~$"):
                    content = get_file_text(file_path, simple_ppt=True)
                    filename = os.path.basename(file_path)
                    content += f"\n\n{filename}"

                    if content:  # 内容が取得できた場合のみインデックス作成
                        writer.add_document(path=str(file_path), content=content)

                # 進捗バーの更新
                progress = (i + 1) / total_files if total_files > 0 else 1
                progress_bar.progress(
                    min(progress, 1),
                    text=f"インデックス作成中... {i + 1} / {total_files} : {file_path}",
                )

            writer.commit()
            progress_bar.empty()  # 完了後に進捗バーを消去

        # Indexが空の場合にはindexを削除しておく
        if ix.doc_count() == 0:
            ix.close()
            shutil.rmtree(index_dir)
            return "No files were indexed in the specified directory."

    # 検索の実行
    ix = index.open_dir(index_dir)
    with ix.searcher() as searcher:
        found_matches: Dict[Path, Set[str]] = {}

        # 各キーワードで個別に検索
        for keyword in keywords:
            query = QueryParser("content", ix.schema).parse(keyword)
            results = searcher.search(query)

            # 検索結果のパスとキーワードを追加
            for result in results:
                path = Path(result["path"])
                if path not in found_matches:
                    found_matches[path] = set()
                found_matches[path].add(keyword)

    # 結果を文字列形式に変換
    result_lines = []
    for path, matched_keywords in found_matches.items():
        keywords_str = ", ".join(matched_keywords)
        result_lines.append(f"File: {path}\nMatched keywords: {keywords_str}\n")

    if not result_lines:
        return "No matches found."

    return "\n".join(result_lines)


def _get_directory_tree(dir_path: str, depth: int, current_depth: int = 0, prefix: str = "") -> str:
    """再帰的にディレクトリ構造を取得する"""
    if current_depth > depth:
        return ""

    if not os.path.exists(dir_path):
        return ""

    # 対象とするファイル拡張子
    target_extensions = {".pdf", ".docx", ".pptx", ".csv", ".xls", ".xlsx", ".doc", ".py"}

    result = []

    try:
        entries = sorted(os.scandir(dir_path), key=lambda e: (not e.is_dir(), e.name.lower()))

        for i, entry in enumerate(entries):
            is_last = i == len(entries) - 1
            current_prefix = "└── " if is_last else "├── "
            next_prefix = "    " if is_last else "│   "

            if entry.is_dir() and not entry.name.startswith("."):
                result.append(f"{prefix}{current_prefix}{entry.name}/")
                subtree = _get_directory_tree(entry.path, depth, current_depth + 1, prefix + next_prefix)
                if subtree:
                    result.append(subtree)
            elif entry.is_file():
                # ファイル拡張子の確認
                _, ext = os.path.splitext(entry.name.lower())
                if ext in target_extensions:
                    result.append(f"{prefix}{current_prefix}{entry.name}")

        return "\n".join(result)
    except PermissionError:
        return f"{prefix}[Permission Denied]"


def get_directory_structure(target_dir: str | list[str], depth: int = 2) -> str:
    """ディレクトリ構造を取得する"""
    if isinstance(target_dir, list):
        results = []
        for dir_path in target_dir:
            dir_path = os.path.expanduser(dir_path)
            if not os.path.exists(dir_path):
                results.append(f"指定されたディレクトリ '{dir_path}' は存在しません。")
                continue
            results.append(dir_path)
            results.append(_get_directory_tree(dir_path, depth))
        return "\n\n".join(results)

    # 単一ディレクトリの場合
    target_dir = os.path.expanduser(target_dir)
    if not os.path.exists(target_dir):
        return f"指定されたディレクトリ '{target_dir}' は存在しません。"
    return f"{target_dir}\n{_get_directory_tree(target_dir, depth)}"
