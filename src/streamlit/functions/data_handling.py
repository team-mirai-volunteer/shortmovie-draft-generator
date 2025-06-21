import fitz
import pandas as pd
from docx import Document
from typing import List

import streamlit as st


@st.cache_data()
def extract_text_from_pdf(pdf_file):
    """
    PDFファイルから全体のテキストを抽出する関数

    Args:
        pdf_file: PDFファイルオブジェクト

    Returns:
        str: 抽出された全体のテキスト
    """
    # ページごとのテキストを取得して結合
    pages_text = extract_text_from_pdf_pages(pdf_file)
    return "\n\n".join(pages_text)


@st.cache_data()
def extract_text_from_pdf_pages(pdf_file) -> List[str]:
    """
    PDFファイルから各ページのテキストを抽出し、ページごとのリストとして返す関数

    Args:
        pdf_file: PDFファイルオブジェクト

    Returns:
        List[str]: 各ページのテキストを含むリスト
    """
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    pages_text = []

    for page in doc:
        page_text = ""

        # テーブルの抽出
        tabs = page.find_tables()

        if not tabs.tables:
            # テーブルがない場合は通常のテキスト抽出
            page_text = page.get_text().strip()
        else:
            # テーブルがある場合は、テーブル以外のテキストを抽出し、テーブルを挿入
            last_end = 0

            for table in tabs.tables:
                # テーブル前のテキストを追加
                page_text += (
                    page.get_text(
                        "text",
                        clip=fitz.Rect(0, last_end, page.rect.width, table.bbox[1]),
                    ).strip()
                    + "\n\n"
                )

                # テーブルをマークダウン形式で追加
                table_data = table.extract()

                if len(table_data) >= 3:
                    # 最初の行がタイトル行かどうかを判定
                    # （最初のセル以外が空、またはNoneの場合）
                    first_row = table_data[0]
                    is_title_row = False

                    if len(first_row) > 1:
                        # 最初のセル以外がすべて空またはNoneの場合、タイトル行とみなす
                        non_empty_cells = sum(1 for cell in first_row[1:] if cell and str(cell).strip())
                        if non_empty_cells == 0 and first_row[0] and str(first_row[0]).strip():
                            is_title_row = True

                    if is_title_row:
                        # タイトル行を別途追加し、実際のヘッダーは2行目を使用
                        title = str(first_row[0]).strip()
                        page_text += f"**{title}**\n\n"

                        # 2行目をヘッダー、3行目以降をデータとして使用
                        if len(table_data) > 2:
                            df = pd.DataFrame(table_data[2:], columns=table_data[1])
                        else:
                            # データ行がない場合は、2行目だけを表示
                            df = pd.DataFrame([table_data[1]])
                    else:
                        # 通常のテーブル処理
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                else:
                    # 行数が少ない場合は通常の処理
                    if len(table_data) > 1:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                    else:
                        df = pd.DataFrame(table_data)

                markdown_table = df.to_markdown(index=False)
                page_text += markdown_table + "\n\n"

                last_end = table.bbox[3]

            # 最後のテーブル以降のテキストを追加
            page_text += page.get_text(
                "text",
                clip=fitz.Rect(0, last_end, page.rect.width, page.rect.height),
            ).strip()

        pages_text.append(page_text.strip())

    return pages_text


@st.cache_data()
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    table_index = 0  # テーブルのインデックスを追跡

    def process_paragraph_element(element):
        """段落要素からテキストを抽出し、フィールドコードも処理する"""
        para_text = ""

        # 段落内の各run要素を処理
        for run_element in element.findall(".//w:r", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            # フィールドコードを検出
            field_char_elements = run_element.findall(".//w:fldChar", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if field_char_elements:
                continue  # フィールド制御文字はスキップ

            # インストラクションテキスト（フィールドコード）を処理
            instr_text_elements = run_element.findall(".//w:instrText", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            if instr_text_elements:
                # フィールドコードの実際の値を取得（必要に応じて）
                for instr in instr_text_elements:
                    instr_content = instr.text if instr.text else ""
                    # フィールドコードの処理（例：テーブル番号など）
                    if "SEQ Table" in instr_content or "STYLEREF" in instr_content:
                        # このフィールドは最終的なドキュメントでレンダリングされる値に置き換えられるので、
                        # ここでは処理せず、別途処理するか、結果としてレンダリングされた値を使用
                        pass
                continue

            # 通常のテキスト
            text_elements = run_element.findall(".//w:t", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            for t in text_elements:
                if t.text:
                    para_text += t.text

        return para_text.strip()

    for element in doc.element.body:
        if element.tag.endswith("p"):  # 段落の場合
            paragraph_text = process_paragraph_element(element)
            if paragraph_text:
                text += paragraph_text + "\n\n"

        elif element.tag.endswith("tbl"):  # テーブルの場合
            # 現在のインデックスのテーブルを取得
            table_element = doc.tables[table_index]
            rows = []

            for row in table_element.rows:
                # セル内のテキストを処理（各セル内の段落も適切に処理）
                row_data = []
                for cell in row.cells:
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        # 段落内のテキストを処理（フィールドコードを含む）
                        para_element = paragraph._element
                        para_text = process_paragraph_element(para_element)
                        if para_text:
                            cell_text += para_text + " "

                    # セル内の改行を<br>に置換
                    row_data.append(cell_text.strip().replace("\n", "<br>"))

                rows.append(row_data)

            if rows:
                max_cols = max(len(row) for row in rows)
                padded_rows = [row + [""] * (max_cols - len(row)) for row in rows]

                if len(padded_rows) >= 3:
                    # 最初の行がタイトル行かどうかを判定
                    first_row = padded_rows[0]
                    is_title_row = False

                    if len(first_row) > 1:
                        # 最初のセル以外がすべて空の場合、タイトル行とみなす
                        non_empty_cells = sum(1 for cell in first_row[1:] if cell and cell.strip())
                        if non_empty_cells == 0 and first_row[0] and first_row[0].strip():
                            is_title_row = True

                    if is_title_row:
                        # タイトル行を別途追加し、実際のヘッダーは2行目を使用
                        title = first_row[0].strip()
                        text += f"**{title}**\n\n"

                        # 2行目をヘッダー、3行目以降をデータとして使用
                        headers = padded_rows[1]
                        if all(not header for header in headers):
                            headers = [f"Column {i+1}" for i in range(len(headers))]

                        if len(padded_rows) > 2:
                            df = pd.DataFrame(padded_rows[2:], columns=headers)
                        else:
                            # データ行がない場合は、ヘッダーのみを表示
                            df = pd.DataFrame([headers])

                        markdown_table = df.to_markdown(index=False)
                        text += markdown_table + "\n\n"
                    else:
                        # 通常のテーブル処理
                        headers = padded_rows[0]
                        if all(not header for header in headers):
                            headers = [f"Column {i+1}" for i in range(len(headers))]

                        df = pd.DataFrame(padded_rows[1:], columns=headers)
                        markdown_table = df.to_markdown(index=False)
                        text += markdown_table + "\n\n"
                elif len(padded_rows) > 1:
                    # 行数が少ない場合は通常の処理
                    headers = padded_rows[0]
                    if all(not header for header in headers):
                        headers = [f"Column {i+1}" for i in range(len(headers))]

                    df = pd.DataFrame(padded_rows[1:], columns=headers)
                    markdown_table = df.to_markdown(index=False)
                    text += markdown_table + "\n\n"
                else:
                    # 1行のみの場合
                    df = pd.DataFrame(padded_rows)
                    markdown_table = df.to_markdown(index=False)
                    text += markdown_table + "\n\n"

            table_index += 1  # 次のテーブルのために増加

    return text.strip()


def read_document(file):
    """
    アップロードされたファイルからテキストを抽出する関数

    Args:
        file: アップロードされたファイルオブジェクト
            - PDFファイル (.pdf)
            - Wordファイル (.docx)
            に対応

    Returns:
        str: 抽出されたテキスト

    処理の説明:
    1. file.typeでファイルのMIMEタイプを確認します
       - "application/pdf": PDFファイル
       - "application/vnd.openxmlformats-officedocument.wordprocessingml.document": Wordファイル

    2. ファイルタイプに応じて適切な抽出関数を呼び出します
       - PDFの場合: extract_text_from_pdf()を使用してテキストを抽出
       - Wordの場合: extract_text_from_docx()を使用してテキストを抽出

    3. 抽出されたテキストを文字列として返します
    """
    if file.type == "application/pdf":
        return extract_text_from_pdf(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file)
